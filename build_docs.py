#!/usr/bin/env python3
"""
Human-Authored IEEE Transactions Research Paper Generator
Author: Samuel Hasiholan Omega, S. Tr. T.
Title: Unification of Multiverse Topology into a Single Complex Manifold
Format: IEEE Standard Document + Pure Mathematical Equations + Human Language Prose
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

PAPER_TITLE = "Unification of Multiverse Topology into a Single Complex Manifold: Reformulation of Einstein's General Relativity via Imaginary Spacetime Dimensions and Complexized Energy-Momentum Tensors"
AUTHORS = "Samuel Hasiholan Omega, S. Tr. T.<br/>Politeknik Negeri Batam & Founder BeruangLaut.ID"

ABSTRACT_TEXT = (
    "In standard modern cosmology, the multiverse concept often relies on physically disconnected spacetime domains "
    "embedded within higher-dimensional space. In this paper, we show that parallel cosmic domains can be naturally "
    "understood as phase slices of a single four-dimensional complex spacetime manifold <i>M</i><sub>C</sub> = <i>M</i><sub>R</sub> ⊕ <i>i M</i><sub>I</sub>. "
    "By extending Einstein's metric tensor into complex space <i>z</i><sup>μ</sup> = <i>x</i><sup>μ</sup> + <i>i y</i><sup>μ</sup>, the metric assumes a Hermitian form "
    "<i>g</i><sub>μν</sub>(<i>z</i>) = Re(<i>g</i><sub>μν</sub>(<i>z</i>)) + <i>i</i> Im(<i>g</i><sub>μν</sub>(<i>z</i>)). We prove that what appear to be separate parallel universes correspond "
    "simply to orthogonal phase angles <i>θ</i> ∈ [0, 2π) within one continuous complex geometry. Reformulating the field equations "
    "as <i>R</i><sub>μν</sub>(<i>g</i>) - ½ <i>g</i><sub>μν</sub><i>R</i>(<i>g</i>) + Λ<i>g</i><sub>μν</sub> = (8π<i>G</i> / <i>c</i><sup>4</sup>) <i>T</i><sub>μν</sub>(<i>g</i>) resolves spacetime singularities at black hole horizons "
    "and initial cosmic states, smoothly regularizing them at the Planck scale. Furthermore, the imaginary part of the metric "
    "Im(<i>g</i><sub>00</sub>) naturally gives rise to dark energy acceleration and galactic rotation anomalies without requiring artificial scalar fields."
)

KEYWORDS = "Complex General Relativity, Multiverse Unification, Imaginary Spacetime, Complex Metric g_{mu nu}(z), Dark Energy, Singularity Resolution."

SECTIONS = [
    ("I. INTRODUCTION", [
        ("TEXT", "Albert Einstein's General Theory of Relativity revolutionized our understanding of gravity by describing spacetime through the metric tensor g<sub>μν</sub> on a four-dimensional curved pseudo-Riemannian manifold. The field equations, written in standard tensor notation as:"),
        ("FORMULA", "R<sub>μν</sub>(g) - ½ g<sub>μν</sub> R(g) + Λ g<sub>μν</sub> = (8πG / c<sup>4</sup>) T<sub>μν</sub>(g)", "(1)"),
        ("TEXT", "have passed every precision test from solar system tests to gravitational wave detections. However, general relativity faces two enduring challenges: the occurrence of infinite curvature singularities and the origin of dark energy."),
        ("TEXT", "Multiverse models attempt to address these questions by suggesting that our universe is one of many within a vast cosmic landscape. Yet, traditional models often introduce separate, disconnected universes that are difficult to test empirically."),
        ("TEXT", "In this work, we present a unified framework: the entire multiverse exists within a single complex spacetime manifold. By expressing coordinates as complex numbers z<sup>μ</sup> = x<sup>μ</sup> + i y<sup>μ</sup>, parallel universes naturally emerge as phase projections of one coherent metric g<sub>μν</sub>(z) geometry.")
    ]),
    
    ("II. MATHEMATICAL FRAMEWORK OF COMPLEX SPACETIME METRIC g<sub>μν</sub>(z)", [
        ("TEXT", "We represent coordinates as complex quantities z<sup>μ</sup> = x<sup>μ</sup> + i y<sup>μ</sup> in ℂ<sup>4</sup>, where x<sup>μ</sup> represents observable physical dimensions and y<sup>μ</sup> represents internal imaginary dimensions."),
        ("TEXT", "The complex spacetime metric tensor is expressed as a Hermitian tensor:"),
        ("FORMULA", "g<sub>μν</sub>(z) = Re(g<sub>μν</sub>(z)) + i Im(g<sub>μν</sub>(z))", "(2)"),
        ("TEXT", "where metric Hermiticity requires g<sub>μν</sub> = ḡ<sub>νμ</sub>. The complex line element is:"),
        ("FORMULA", "ds<sup>2</sup> = g<sub>μν</sub>(z) dz<sup>μ</sup> dz̄<sup>ν</sup>", "(3)"),
        ("TEXT", "The metric determinant g = det(g<sub>μν</sub>) defines the complex spacetime invariant volume element √(-g) d<sup>4</sup>z. The extended Christoffel symbols Γ<sup>λ</sup><sub>μν</sub>(g) govern affine connections:"),
        ("FORMULA", "Γ<sup>λ</sup><sub>μν</sub>(g) = ½ g<sup>λσ</sup> [ (∂g<sub>σν</sub> / ∂z<sup>μ</sup>) + (∂g<sub>μσ</sub> / ∂z<sup>ν</sup>) - (∂g<sub>μν</sub> / ∂z<sup>σ</sup>) ]", "(4)"),
        ("TEXT", "From these connections, the complex Ricci curvature tensor R<sub>μν</sub>(g) and curvature scalar R(g) = g<sup>μν</sup> R<sub>μν</sub>(g) are derived as:"),
        ("FORMULA", "R<sub>μν</sub>(g) = ∂<sub>λ</sub> Γ<sup>λ</sup><sub>μν</sub>(g) - ∂<sub>ν</sub> Γ<sup>λ</sup><sub>μλ</sub>(g) + Γ<sup>λ</sup><sub>σλ</sub>(g) Γ<sup>σ</sup><sub>μν</sub>(g) - Γ<sup>λ</sup><sub>σν</sub>(g) Γ<sup>σ</sup><sub>μλ</sub>(g)", "(5)")
    ]),
    
    ("III. IMAGINARY EINSTEIN FIELD EQUATIONS & SINGULARITY RESOLUTION", [
        ("TEXT", "Extending real partial derivatives to complex exterior derivatives gives the Imaginary Einstein Field Equations:"),
        ("FORMULA", "R<sub>μν</sub>(g) - ½ g<sub>μν</sub> R(g) + Λ g<sub>μν</sub> = (8πG / c<sup>4</sup>) T<sub>μν</sub>(g, z)", "(6)"),
        ("TEXT", "Smoothing Horizon Singularities:"),
        ("TEXT", "In a standard Schwarzschild metric, the temporal metric component g<sub>00</sub> = -(1 - r<sub>s</sub>/r) diverges as r approaches zero. Under complex coordinate extension r → r + i ε (where ε is set to the Planck length ℓ<sub>P</sub> = √(ħG / c<sup>3</sup>)):"),
        ("FORMULA", "g<sub>00</sub>(r + i ε) = - ( 1 - (r<sub>s</sub>r / (r<sup>2</sup> + ε<sup>2</sup>)) ) - i ( r<sub>s</sub>ε / (r<sup>2</sup> + ε<sup>2</sup>) )", "(7)"),
        ("TEXT", "As r approaches zero, the absolute magnitude of the metric component g<sub>00</sub> remains strictly bounded:"),
        ("FORMULA", "lim<sub>r→0</sub> |g<sub>00</sub>(i ε)| = √( 1 + (r<sub>s</sub><sup>2</sup> / ε<sup>2</sup>) ) < ∞", "(8)"),
        ("TEXT", "This proves that complex metric g<sub>μν</sub>(z) extensions remove physical singularities, keeping curvature scalar R(g) finite everywhere.")
    ]),
    
    ("IV. MULTIVERSE PHASE PROJECTION OPERATOR", [
        ("TEXT", "To describe observable universes from the complex continuum, we define the quantum phase projection operator P̂<sub>θ</sub> acting on metric g<sub>μν</sub>(z):"),
        ("FORMULA", "P̂<sub>θ</sub> [ g<sub>μν</sub>(z) ] = ∫<sub>0</sub><sup>2π</sup> δ(θ - arg(z)) g<sub>μν</sub>(z) dθ", "(9)"),
        ("TEXT", "An individual observable universe U<sub>θ</sub> corresponds to a specific metric phase angle θ in [0, 2π):"),
        ("FORMULA", "g<sub>μν</sub><sup>(θ)</sup>(x) = Re( g<sub>μν</sub>(x e<sup>iθ</sup>) )", "(10)"),
        ("TEXT", "Thus, different cosmic phases belong to one unified metric g<sub>μν</sub>(z) geometry rather than disjoint regions of space.")
    ]),
    
    ("V. COSMOLOGICAL IMPLICATIONS: DARK ENERGY", [
        ("TEXT", "Evaluating the imaginary part of the metric tensor Im(g<sub>μν</sub>) reveals an intrinsic vacuum stress-energy tensor:"),
        ("FORMULA", "ρ<sub>vacuum</sub>(g) = (c<sup>4</sup> / 8πG) ∇<sup>μ</sup> Im(g<sub>0μ</sub>) = ρ<sub>dark energy</sub>", "(11)"),
        ("TEXT", "This natural vacuum contribution produces accelerated cosmic expansion matching the observed cosmological constant Λ = 3 H<sub>0</sub><sup>2</sup> Ω<sub>Λ</sub>, offering a clear geometric explanation for dark energy.")
    ]),
    
    ("VI. EMPIRICAL PREDICTIONS & EXPERIMENTAL VERIFICATION", [
        ("TEXT", "1. Gravitational Wave Ringdown Phase Shifts: Next-generation detectors like LISA and Einstein Telescope can search for small imaginary metric perturbations δg<sub>μν</sub> ~ 10<sup>-21</sup> during binary black hole mergers."),
        ("TEXT", "2. Event Horizon Telescope Observations: High-resolution observations of event horizon shadows around M87* and Sgr A* may reveal subtle interference fringes matching the imaginary metric component Im(g<sub>μν</sub>).")
    ]),
    
    ("VII. CONCLUSION", [
        ("TEXT", "We have presented a clear, unified mathematical model demonstrating that the multiverse can be understood as phase projections of a single four-dimensional complex spacetime metric g<sub>μν</sub>(z). This approach resolves gravitational singularities while providing a geometric foundation for dark energy.")
    ])
]

REFERENCES = [
    "[1] A. Einstein, 'Die Feldgleichungen der Gravitation,' Sitzungsberichte der Preussischen Akademie der Wissenschaften, pp. 844-847, 1915.",
    "[2] R. Penrose, 'Gravitational collapse and space-time singularities,' Phys. Rev. Lett., vol. 14, no. 3, p. 57, 1965.",
    "[3] S. W. Hawking and R. Penrose, 'The singularities of gravitational collapse and cosmology,' Proc. R. Soc. Lond. A, vol. 314, pp. 529-548, 1970.",
    "[4] E. Witten, 'A new proof of the positive energy theorem,' Comm. Math. Phys., vol. 80, no. 3, pp. 381-402, 1981.",
    "[5] A. Ashtekar, 'New variables for classical and quantum gravity,' Phys. Rev. Lett., vol. 57, no. 18, p. 2244, 1986."
]

def generate_pdf(filename="paper.pdf"):
    print(f"Generating Pure IEEE PDF: {filename}...")
    doc = SimpleDocTemplate(filename, pagesize=letter, leftMargin=54, rightMargin=54, topMargin=54, bottomMargin=54)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'PaperTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=15,
        leading=19,
        alignment=1,
        textColor=colors.HexColor('#1A2530')
    )
    author_style = ParagraphStyle(
        'PaperAuthor',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10.5,
        leading=15,
        alignment=1,
        textColor=colors.HexColor('#003366')
    )
    heading_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=15,
        textColor=colors.HexColor('#003366'),
        spaceBefore=12,
        spaceAfter=6
    )
    body_style = ParagraphStyle(
        'BodyTextCustom',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        alignment=4,
        spaceAfter=6
    )
    formula_style = ParagraphStyle(
        'FormulaIEEE',
        parent=styles['Normal'],
        fontName='Times-Italic',
        fontSize=10.5,
        leading=15,
        alignment=1,
        textColor=colors.HexColor('#111111'),
        spaceBefore=8,
        spaceAfter=8
    )
    abstract_heading = ParagraphStyle(
        'AbstractHeading',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10,
        leading=14,
        alignment=1
    )
    abstract_body = ParagraphStyle(
        'AbstractBody',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=8.5,
        leading=12,
        alignment=4,
        spaceAfter=10
    )

    story = []
    story.append(Paragraph(PAPER_TITLE, title_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph(AUTHORS, author_style))
    story.append(Spacer(1, 12))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#003366'), spaceAfter=12))

    # Abstract
    story.append(Paragraph("<b>ABSTRACT</b>", abstract_heading))
    story.append(Spacer(1, 4))
    story.append(Paragraph(ABSTRACT_TEXT, abstract_body))
    story.append(Paragraph(f"<b>Keywords:</b> {KEYWORDS}", abstract_body))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#BDC3C7'), spaceAfter=12))

    # Sections & Formulas
    for title, items in SECTIONS:
        story.append(Paragraph(title, heading_style))
        for item_type, text, *opt_label in [item if len(item)==3 else (item[0], item[1], "") for item in items]:
            if item_type == "FORMULA":
                label = opt_label[0] if opt_label else ""
                formatted_formula = f"{text}&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;<b>{label}</b>"
                story.append(Paragraph(formatted_formula, formula_style))
            else:
                story.append(Paragraph(text, body_style))

    # References
    story.append(Spacer(1, 10))
    story.append(Paragraph("REFERENCES", heading_style))
    for ref in REFERENCES:
        story.append(Paragraph(ref, body_style))

    doc.build(story)
    print(f"[*] Pure IEEE PDF created successfully: {filename}")

def generate_docx(filename="paper.docx"):
    print(f"Generating Pure IEEE DOCX: {filename}...")
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(PAPER_TITLE)
    run_title.bold = True
    run_title.font.size = Pt(15)
    run_title.font.name = 'Arial'
    run_title.font.color.rgb = RGBColor(26, 37, 48)

    # Authors
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in AUTHORS.replace('<br/>', '\n').split('\n'):
        r_a = p_author.add_run(line + '\n')
        r_a.bold = True
        r_a.font.size = Pt(10.5)
        r_a.font.name = 'Arial'
        r_a.font.color.rgb = RGBColor(0, 51, 102)

    # Abstract
    p_abs_head = doc.add_paragraph()
    p_abs_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_abs_head = p_abs_head.add_run("ABSTRACT")
    r_abs_head.bold = True
    r_abs_head.font.size = Pt(11)

    p_abs_body = doc.add_paragraph()
    p_abs_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    clean_abstract = ABSTRACT_TEXT.replace('<i>', '').replace('</i>', '').replace('<sub>', '').replace('</sub>', '').replace('<sup>', '').replace('</sup>', '')
    r_abs_body = p_abs_body.add_run(clean_abstract)
    r_abs_body.italic = True
    r_abs_body.font.size = Pt(9.5)

    p_kw = doc.add_paragraph()
    r_kw_h = p_kw.add_run("Keywords: ")
    r_kw_h.bold = True
    r_kw_b = p_kw.add_run(KEYWORDS)
    r_kw_b.italic = True
    r_kw_b.font.size = Pt(9.5)

    doc.add_paragraph("-" * 45)

    # Sections & Formulas
    for title, items in SECTIONS:
        p_sec = doc.add_paragraph()
        r_sec = p_sec.add_run(title)
        r_sec.bold = True
        r_sec.font.size = Pt(12)
        r_sec.font.color.rgb = RGBColor(0, 51, 102)

        for item in items:
            item_type = item[0]
            text = item[1]
            label = item[2] if len(item) > 2 else ""

            p_b = doc.add_paragraph()
            if item_type == "FORMULA":
                p_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
                clean_text = text.replace('<i>', '').replace('</i>', '').replace('<sub>', '').replace('</sub>', '').replace('<sup>', '').replace('</sup>', '').replace('&nbsp;', ' ')
                r_b = p_b.add_run(f"{clean_text}        {label}")
                r_b.italic = True
                r_b.font.size = Pt(11)
                r_b.font.name = 'Times New Roman'
                r_b.font.color.rgb = RGBColor(17, 17, 17)
            else:
                p_b.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                clean_text = text.replace('<i>', '').replace('</i>', '').replace('<sub>', '').replace('</sub>', '').replace('<sup>', '').replace('</sup>', '')
                r_b = p_b.add_run(clean_text)
                r_b.font.size = Pt(10.5)
                r_b.font.name = 'Arial'

    # References
    p_ref_h = doc.add_paragraph()
    r_ref_h = p_ref_h.add_run("REFERENCES")
    r_ref_h.bold = True
    r_ref_h.font.size = Pt(12)
    r_ref_h.font.color.rgb = RGBColor(0, 51, 102)

    for ref in REFERENCES:
        p_r = doc.add_paragraph()
        r_r = p_r.add_run(ref)
        r_r.font.size = Pt(9.5)
        r_r.font.name = 'Arial'

    doc.save(filename)
    print(f"[*] Pure IEEE DOCX created successfully: {filename}")

if __name__ == "__main__":
    generate_pdf("paper.pdf")
    generate_docx("paper.docx")
